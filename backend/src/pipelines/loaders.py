import re
from pathlib import Path
from typing import Iterable

from langchain_core.documents import Document as LCDocument
from langchain_community.document_loaders import PyPDFLoader


# ─── Table-aware text preprocessing ─────────────────────────────────────────
# Tables in PDFs/DOCX often get chunked into orphaned rows that lose their
# column headers. "JWT (httpOnly cookies)" becomes meaningless without knowing
# it's in the "Authentication" row under "Technology Stack". Fix: detect
# table-like structures and serialize each row as a natural-language sentence
# BEFORE the text splitter runs.
#
# Pattern detected: lines with | separators (markdown tables) or consistent
# tab/multi-space alignment. We convert:
#   | Layer | Technology |
#   | Authentication | JWT (httpOnly cookies) |
# Into:
#   "Layer: Authentication. Technology: JWT (httpOnly cookies)."

_TABLE_ROW_RE = re.compile(r"^\s*\|(.+)\|\s*$")
_TABLE_SEP_RE = re.compile(r"^\s*\|[\s\-:|]+\|\s*$")  # separator like |---|---|


def _serialize_tables(text: str) -> str:
    """Convert markdown-style tables into natural-language sentences so
    table cells retain their column-header context after chunking.

    Non-table text passes through unchanged. Multiple tables in one
    document are handled independently.
    """
    lines = text.split("\n")
    result: list[str] = []
    i = 0

    while i < len(lines):
        line = lines[i]

        # Detect table start: a line with | separators.
        if _TABLE_ROW_RE.match(line):
            # Collect the full table.
            table_lines: list[str] = []
            while i < len(lines) and (_TABLE_ROW_RE.match(lines[i]) or _TABLE_SEP_RE.match(lines[i])):
                if not _TABLE_SEP_RE.match(lines[i]):  # skip separator rows
                    table_lines.append(lines[i])
                i += 1

            if len(table_lines) >= 2:
                # First row = headers, rest = data rows.
                headers = [
                    h.strip() for h in table_lines[0].split("|") if h.strip()
                ]
                for row_line in table_lines[1:]:
                    cells = [c.strip() for c in row_line.split("|") if c.strip()]
                    # Build sentence: "Header1: Cell1. Header2: Cell2."
                    parts = []
                    for j, cell in enumerate(cells):
                        header = headers[j] if j < len(headers) else f"Column {j+1}"
                        parts.append(f"{header}: {cell}")
                    sentence = ". ".join(parts) + "."
                    result.append(sentence)
                result.append("")  # blank line after serialized table
            else:
                # Single-row "table" — just pass through.
                for tl in table_lines:
                    result.append(tl)
                i += 1
        else:
            result.append(line)
            i += 1

    return "\n".join(result)


def _preprocess_text(text: str) -> str:
    """Run all text preprocessing steps before chunking.
    Currently: table serialization. Extensible for future transforms.
    """
    return _serialize_tables(text)


# Tier 2.1 — minimum chars per page below which we suspect a scanned PDF.
# PyPDFLoader returns empty/near-empty page_content for image-only pages;
# we kick those over to Tesseract via pdf2image for OCR.
_OCR_MIN_CHARS_PER_PAGE = 40


def _ocr_pdf(path: Path) -> list[LCDocument]:
    """Run Tesseract OCR on each page of a PDF. Used as a fallback for
    scanned/image-only PDFs where PyPDFLoader returns empty text. Slow
    (~1-3s per page on CPU) but unavoidable for non-OCR'd scans.

    Requires system: tesseract + poppler (brew install tesseract poppler).
    Requires Python: pytesseract + pdf2image.
    """
    try:
        import pytesseract
        from pdf2image import convert_from_path
    except ImportError:
        return []

    try:
        # 200 DPI is the sweet spot — high enough for clean OCR on body
        # text, low enough to keep each page render under ~1.5s on CPU.
        images = convert_from_path(str(path), dpi=200)
    except Exception:
        return []

    out: list[LCDocument] = []
    for i, img in enumerate(images):
        try:
            text = pytesseract.image_to_string(img, lang="eng").strip()
        except Exception:
            text = ""
        if text:
            out.append(
                LCDocument(
                    page_content=text,
                    metadata={"page": i + 1, "source": "ocr"},
                )
            )
    return out


def _extract_pdf_tables(path: Path) -> dict[int, list[str]]:
    """Extract tables from a PDF using pdfplumber and serialize each row
    as 'Header: Value' sentences. Returns {page_number: [serialized_rows]}.

    Falls back to empty dict if pdfplumber is unavailable.
    """
    try:
        import pdfplumber
    except ImportError:
        return {}

    tables_by_page: dict[int, list[str]] = {}
    try:
        with pdfplumber.open(str(path)) as pdf:
            for page_num, page in enumerate(pdf.pages):
                page_tables = page.extract_tables()
                if not page_tables:
                    continue
                serialized: list[str] = []
                for table in page_tables:
                    if len(table) < 2:
                        continue
                    # First row = headers
                    headers = [(h or "").strip() or f"Column {j+1}"
                               for j, h in enumerate(table[0])]
                    for row in table[1:]:
                        cells = [(c or "").strip() for c in row]
                        if not any(cells):
                            continue
                        parts = []
                        for j, cell_val in enumerate(cells):
                            if not cell_val:
                                continue
                            header = headers[j] if j < len(headers) else f"Column {j+1}"
                            parts.append(f"{header}: {cell_val}")
                        if parts:
                            serialized.append(". ".join(parts) + ".")
                if serialized:
                    tables_by_page[page_num] = serialized
    except Exception:
        pass
    return tables_by_page


def _load_pdf(path: Path) -> list[LCDocument]:
    docs = PyPDFLoader(str(path)).load()
    for d in docs:
        if "page" not in d.metadata:
            d.metadata["page"] = d.metadata.get("page_number", 0)

    # Tier 2.1 — OCR fallback for scanned PDFs. If PyPDFLoader produced
    # no docs OR the average page has < _OCR_MIN_CHARS_PER_PAGE of
    # extracted text, the PDF is almost certainly an image-only scan.
    # Re-process via Tesseract.
    text_total = sum(len((d.page_content or "").strip()) for d in docs)
    needs_ocr = (
        not docs
        or (docs and text_total / max(len(docs), 1) < _OCR_MIN_CHARS_PER_PAGE)
    )
    if needs_ocr:
        ocr_docs = _ocr_pdf(path)
        if ocr_docs:
            docs = ocr_docs

    # Extract structured tables via pdfplumber and append serialized rows
    # to the page content. This preserves column headers on every row so
    # table data survives chunking — same fix as DOCX/Excel.
    pdf_tables = _extract_pdf_tables(path)
    for d in docs:
        page_idx = d.metadata.get("page", 0)
        # PyPDFLoader uses 0-based page index; pdfplumber also 0-based.
        if page_idx in pdf_tables:
            table_block = "\n".join(pdf_tables[page_idx])
            d.page_content = (d.page_content or "") + "\n\n" + table_block

    # Table-aware preprocessing: serialize markdown-style tables into
    # natural-language sentences (catches any tables pdfplumber missed).
    for d in docs:
        d.page_content = _preprocess_text(d.page_content or "")
    return docs


def _docx_page_count(path: Path) -> int:
    """Read the real page count from docProps/app.xml inside the DOCX zip.
    Word writes <Pages>N</Pages> on save; returns 0 if absent (older/odd docs)."""
    import xml.etree.ElementTree as ET
    import zipfile

    try:
        with zipfile.ZipFile(path) as z:
            with z.open("docProps/app.xml") as f:
                root = ET.parse(f).getroot()
        # Namespace-agnostic lookup — property lives under the extended-properties ns.
        for child in root.iter():
            tag = child.tag.split("}", 1)[-1]
            if tag == "Pages" and (child.text or "").strip().isdigit():
                return int(child.text.strip())
    except Exception:
        pass
    return 0


def _extract_docx_with_tables(path: Path) -> str:
    """Extract DOCX content using python-docx, serializing tables with
    column headers on every row so table structure survives chunking.

    Falls back to docx2txt if python-docx fails.
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        import docx2txt
        return (docx2txt.process(str(path)) or "").strip()

    try:
        doc = DocxDocument(str(path))
    except Exception:
        import docx2txt
        return (docx2txt.process(str(path)) or "").strip()

    # Build a content-order list from the document body. python-docx
    # exposes body.iter_inner_content() in newer versions but we use
    # the XML element tree for compatibility with all versions.
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    parts: list[str] = []
    for element in doc.element.body:
        tag = element.tag.split("}")[-1]

        if tag == "p":
            para = Paragraph(element, doc)
            text = para.text.strip()
            if text:
                parts.append(text)

        elif tag == "tbl":
            table = Table(element, doc)
            if len(table.rows) < 2:
                # Single-row table — just dump cell values.
                cells = [c.text.strip() for c in table.rows[0].cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
                continue

            headers = [cell.text.strip() for cell in table.rows[0].cells]
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                row_parts = []
                for j, cell_val in enumerate(cells):
                    if not cell_val:
                        continue
                    header = headers[j] if j < len(headers) else f"Column {j+1}"
                    row_parts.append(f"{header}: {cell_val}")
                if row_parts:
                    parts.append(". ".join(row_parts) + ".")
            parts.append("")  # blank line after table

    return "\n\n".join(parts)


def _load_docx(path: Path) -> list[LCDocument]:
    """Load a .docx as N pseudo-pages (N = real page count when available,
    otherwise estimated from text length at ~3000 chars/page). This matches
    PDF/text behavior — one LCDocument per page, then the standard text
    splitter produces chunks on top. Fixes the old behavior that treated
    each paragraph as a page and produced hundreds of fake pages."""

    text = _preprocess_text(_extract_docx_with_tables(path))
    if not text:
        return [LCDocument(page_content="", metadata={"page": 1})]

    pages_meta = _docx_page_count(path)
    if pages_meta > 0:
        n_pages = pages_meta
    else:
        # Fallback estimate: ~3000 chars per printed page.
        n_pages = max(1, (len(text) + 2999) // 3000)

    # Split the concatenated text into roughly equal pseudo-page spans so page
    # citations are stable and meaningful. Splits prefer paragraph boundaries
    # where possible.
    if n_pages == 1:
        return [LCDocument(page_content=text, metadata={"page": 1})]

    span = max(1, len(text) // n_pages)
    docs: list[LCDocument] = []
    cursor = 0
    for i in range(n_pages):
        end = len(text) if i == n_pages - 1 else cursor + span
        # Nudge the cut to the next paragraph boundary so pages read naturally.
        if end < len(text):
            nl = text.find("\n\n", end)
            if nl != -1 and nl - end < 400:
                end = nl
        page_text = text[cursor:end].strip()
        if page_text:
            docs.append(LCDocument(page_content=page_text, metadata={"page": i + 1}))
        cursor = end
    return docs or [LCDocument(page_content=text, metadata={"page": 1})]


def _load_text(path: Path) -> list[LCDocument]:
    text = _preprocess_text(path.read_text(encoding="utf-8", errors="ignore"))
    # split into ~3000-char pseudo-pages so retrieval scores reference a stable page number
    page_size = 3000
    chunks = [text[i : i + page_size] for i in range(0, len(text), page_size)] or [text]
    return [LCDocument(page_content=c, metadata={"page": i + 1}) for i, c in enumerate(chunks)]


def _compute_column_stats(headers: list[str], data_rows: list[list]) -> str:
    """Pre-compute aggregate statistics for numeric columns.

    Returns a natural-language summary block:
      "Column Statistics Summary:
       Bank Charges: count=76, total=8899.56, average=117.10, min=1.00, max=898.80
       ..."

    This chunk is placed at the start of the document so aggregation
    queries ("total bank charges", "how many transactions") retrieve it
    directly instead of trying to sum individual rows.
    """
    stats_lines = []
    for j, header in enumerate(headers):
        values: list[float] = []
        for row in data_rows:
            if j < len(row):
                raw = row[j]
                if raw is None:
                    continue
                # Try to parse as number
                try:
                    v = float(str(raw).replace(",", "").strip())
                    values.append(v)
                except (ValueError, TypeError):
                    continue
        if len(values) >= 3:  # need at least 3 numeric values to be meaningful
            total = sum(values)
            avg = total / len(values)
            stats_lines.append(
                f"{header}: count={len(values)}, total={total:.2f}, "
                f"average={avg:.2f}, min={min(values):.2f}, max={max(values):.2f}"
            )

    if not stats_lines:
        return ""

    return "Column Statistics Summary (pre-computed from all rows):\n" + "\n".join(stats_lines)


def _build_cross_sheet_summary(wb) -> str:
    """For multi-sheet workbooks where each sheet represents an entity,
    build a cross-sheet comparison table.

    Works generically on any multi-sheet workbook:
    1. Extracts metadata from header rows (rows before the data table)
       — key:value patterns like "Employee: Name" or free text labels
    2. Finds the data header row (first row with 3+ non-empty cells)
    3. Finds statistics/summary rows at the bottom
    4. Builds a one-line-per-sheet summary with all extracted info

    Also detects repeated categorical values across sheets and produces
    group counts (e.g. "By Department: Sales: 40, Finance: 9").
    """
    # Known metadata patterns: "Key: Value" in header rows
    _KV_RE = re.compile(r"([A-Za-z][\w\s]{1,25}):\s*(.+)")
    # Employee-specific pattern (common in HR exports)
    _EMP_RE = re.compile(
        r"Employee:\s*(.+?)(?:,,|\s*,\s*)(\w*)\s*\[(\d+)\].*?Department:\s*(\w[\w\s]*?)(?:\[|$)",
        re.IGNORECASE,
    )

    entries: list[str] = []
    # Track categorical values for group counts
    category_counts: dict[str, dict[str, int]] = {}

    for sheet in wb.worksheets:
        rows = list(sheet.iter_rows(values_only=True))
        if len(rows) < 2:
            continue

        # Find the data header row (first row with 3+ non-empty cells)
        header_idx = 0
        for idx, row in enumerate(rows[:6]):
            non_empty = sum(1 for v in row if v is not None and str(v).strip())
            if non_empty >= 3:
                header_idx = idx
                break

        # Extract metadata from rows above the header
        meta: dict[str, str] = {}
        for row in rows[:header_idx]:
            for cell in row:
                if cell is None:
                    continue
                text = str(cell).strip()
                if not text:
                    continue

                # Try employee-specific pattern first
                m = _EMP_RE.search(text)
                if m:
                    meta["Employee"] = m.group(1).strip().rstrip(",")
                    if m.group(2).strip():
                        meta["Gender"] = m.group(2).strip()
                    meta["ID"] = m.group(3).strip()
                    meta["Department"] = m.group(4).strip()
                    continue

                # Generic key:value patterns
                m2 = _KV_RE.match(text)
                if m2:
                    key = m2.group(1).strip()
                    val = m2.group(2).strip()
                    if val and len(val) < 100:  # skip very long values
                        meta[key] = val
                elif len(text) < 80 and text not in meta.values():
                    # Free text label — use as a title/identifier
                    if "Title" not in meta:
                        meta["Title"] = text

        # Find statistics/summary row at the bottom
        stats: dict[str, str] = {}
        hdrs = [str(h).strip() if h else "" for h in rows[header_idx]]
        for row in reversed(rows):
            first_cell = str(row[0]).strip().lower() if row[0] else ""
            if first_cell in ("statistics", "total", "totals", "summary", "grand total"):
                for j, val in enumerate(row):
                    if val and j < len(hdrs) and hdrs[j] and j > 0:
                        stats[hdrs[j]] = str(val).strip()
                break

        # Build entry — combine metadata + sheet name + stats
        parts = []
        for k, v in meta.items():
            parts.append(f"{k}: {v}")
            # Track categories for group counts
            if k in ("Department", "Gender", "Category", "Region", "Type", "Status",
                      "Team", "Division", "Group", "Class", "Level"):
                if k not in category_counts:
                    category_counts[k] = {}
                category_counts[k][v] = category_counts[k].get(v, 0) + 1

        if not parts:
            # No metadata found — use sheet name as identifier
            parts.append(f"Sheet: {sheet.title}")
        else:
            parts.append(f"Sheet: {sheet.title}")

        for k, v in stats.items():
            parts.append(f"{k}: {v}")

        entries.append(". ".join(parts) + ".")

    if len(entries) < 2:
        return ""

    # Build the summary header with group counts
    summary_parts = [f"Cross-sheet summary ({len(entries)} entities):"]
    for cat_name, counts in sorted(category_counts.items()):
        count_str = ", ".join(f"{k}: {v}" for k, v in sorted(counts.items()))
        summary_parts.append(f"By {cat_name.lower()}: {count_str}.")
    summary_parts.append("")
    summary_parts.extend(entries)

    return "\n".join(summary_parts)


def _load_excel(path: Path) -> list[LCDocument]:
    """Load .xlsx/.xls — serialize each row as 'Header: Value' sentences.

    Multi-sheet workbooks produce one section per sheet. Row 1 of each
    sheet is treated as column headers. Empty rows are skipped.

    Two summary chunks are prepended:
    1. Per-sheet column statistics (totals, counts, averages)
    2. Cross-sheet comparison table (for multi-entity workbooks)
    """
    import openpyxl

    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    docs: list[LCDocument] = []

    # Cross-sheet summary for multi-sheet workbooks (employee comparisons, etc.)
    if len(wb.worksheets) > 3:
        cross_summary = _build_cross_sheet_summary(wb)
        if cross_summary:
            # Split into chunks if very large (66 employees = ~5KB)
            page_size = 3000
            chunks = [cross_summary[i : i + page_size]
                      for i in range(0, len(cross_summary), page_size)]
            for i, chunk in enumerate(chunks):
                docs.append(LCDocument(page_content=chunk, metadata={"page": 0}))

    for sheet in wb.worksheets:
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            continue

        # Find the header row (first row with multiple non-empty cells)
        header_idx = 0
        for idx, row in enumerate(rows[:5]):
            non_empty = sum(1 for v in row if v is not None and str(v).strip())
            if non_empty >= 3:
                header_idx = idx
                break

        headers = [str(h).strip() if h is not None else f"Column {j+1}"
                   for j, h in enumerate(rows[header_idx])]

        # Pre-compute column statistics summary
        data_rows = [list(r) for r in rows[header_idx + 1:]]
        stats_block = _compute_column_stats(headers, data_rows)
        if stats_block:
            # Include sheet metadata (employee name, etc.) in the stats
            meta_lines = []
            for row in rows[:header_idx]:
                for cell in row:
                    if cell and str(cell).strip():
                        meta_lines.append(str(cell).strip())
            meta_prefix = "\n".join(meta_lines[:3]) + "\n\n" if meta_lines else ""

            sheet_prefix = f"Sheet: {sheet.title}\n" if len(wb.worksheets) > 1 else ""
            row_count = len([r for r in data_rows if any(
                v is not None and str(v).strip() for v in r
            )])
            summary = (
                f"{sheet_prefix}{meta_prefix}"
                f"This sheet contains {row_count} data rows with columns: "
                f"{', '.join(headers)}.\n\n{stats_block}"
            )
            docs.append(LCDocument(page_content=summary, metadata={"page": 0}))

        # Serialize individual rows (include metadata rows too)
        parts: list[str] = []
        if len(wb.worksheets) > 1:
            parts.append(f"Sheet: {sheet.title}")

        # Add metadata rows as context
        for row in rows[:header_idx]:
            for cell in row:
                if cell and str(cell).strip():
                    parts.append(str(cell).strip())

        for row in rows[header_idx + 1:]:
            cells = [str(v).strip() if v is not None else "" for v in row]
            if not any(cells):
                continue
            row_parts = []
            for j, cell_val in enumerate(cells):
                if not cell_val:
                    continue
                header = headers[j] if j < len(headers) else f"Column {j+1}"
                row_parts.append(f"{header}: {cell_val}")
            if row_parts:
                parts.append(". ".join(row_parts) + ".")

        text = "\n".join(parts).strip()
        if text:
            page_size = 3000
            chunks = [text[i : i + page_size] for i in range(0, len(text), page_size)]
            for i, chunk in enumerate(chunks):
                docs.append(LCDocument(page_content=chunk, metadata={"page": i + 1}))

    wb.close()
    return docs or [LCDocument(page_content="", metadata={"page": 1})]


def _load_csv(path: Path) -> list[LCDocument]:
    """Load .csv — serialize each row as 'Header: Value' sentences.

    First row is treated as column headers. A pre-computed statistics
    summary is prepended for aggregation queries.
    """
    import csv

    with open(path, newline="", encoding="utf-8", errors="ignore") as f:
        reader = csv.reader(f)
        rows = list(reader)

    if len(rows) < 2:
        text = "\n".join(",".join(r) for r in rows)
        return [LCDocument(page_content=text, metadata={"page": 1})]

    headers = [h.strip() or f"Column {j+1}" for j, h in enumerate(rows[0])]
    data_rows = [row for row in rows[1:] if any(v.strip() for v in row)]
    docs: list[LCDocument] = []

    # Pre-computed stats summary
    stats_block = _compute_column_stats(headers, data_rows)
    if stats_block:
        summary = (
            f"This file contains {len(data_rows)} data rows with columns: "
            f"{', '.join(headers)}.\n\n{stats_block}"
        )
        docs.append(LCDocument(page_content=summary, metadata={"page": 0}))

    # Serialize individual rows
    parts: list[str] = []
    for row in data_rows:
        cells = [v.strip() for v in row]
        row_parts = []
        for j, cell_val in enumerate(cells):
            if not cell_val:
                continue
            header = headers[j] if j < len(headers) else f"Column {j+1}"
            row_parts.append(f"{header}: {cell_val}")
        if row_parts:
            parts.append(". ".join(row_parts) + ".")

    text = "\n".join(parts).strip()
    if text:
        page_size = 3000
        chunks = [text[i : i + page_size] for i in range(0, len(text), page_size)]
        for i, chunk in enumerate(chunks):
            docs.append(LCDocument(page_content=chunk, metadata={"page": i + 1}))

    return docs or [LCDocument(page_content="", metadata={"page": 1})]


SUPPORTED_EXTS = {".pdf", ".docx", ".txt", ".md", ".markdown", ".xlsx", ".xls", ".csv"}


def load_any(path: Path) -> list[LCDocument]:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _load_pdf(path)
    if ext == ".docx":
        return _load_docx(path)
    if ext in {".xlsx", ".xls"}:
        return _load_excel(path)
    if ext == ".csv":
        return _load_csv(path)
    if ext in {".txt", ".md", ".markdown"}:
        return _load_text(path)
    raise ValueError(f"Unsupported file type: {ext}")


def detect_mime(path: Path) -> str:
    ext = path.suffix.lower()
    return {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ".xls": "application/vnd.ms-excel",
        ".csv": "text/csv",
        ".txt": "text/plain",
        ".md": "text/markdown",
        ".markdown": "text/markdown",
    }.get(ext, "application/octet-stream")
